"""Word document renderer using python-docx."""

from typing import Any
from uuid import uuid4
from pathlib import Path
import logging

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class WordRenderer:
    """Renders Word documents from Layout Specification Packages."""

    def __init__(self, output_dir: str | None = None) -> None:
        """
        Initialize Word renderer.

        Args:
            output_dir: Directory to save generated documents (defaults to infrastructure/data/artifacts)
        """
        if output_dir is None:
            # Use absolute path relative to project root
            project_root = Path(__file__).parent.parent.parent.parent.resolve()
            output_dir = project_root / "infrastructure" / "data" / "artifacts"
        self.output_dir = Path(output_dir).resolve()
        self.output_dir.mkdir(parents=True, exist_ok=True)

    async def render(self, lsp: dict[str, Any]) -> str:
        """
        Render Word document from LSP.

        Args:
            lsp: Layout Specification Package

        Returns:
            Artifact ID for the generated document
        """
        # Extract metadata
        metadata = lsp.get("metadata", {})
        structure = lsp.get("structure", [])
        proposal_id = lsp.get("proposal_id", "unknown")
        session_id = lsp.get("session_id", "unknown")

        # Create document
        doc = Document()

        # Set document properties
        title = metadata.get("title", "Untitled Document")
        doc.core_properties.title = title

        logger.info(f"Rendering Word document: {title} with {len(structure)} pages")
        logger.debug(f"Content map has {len(lsp.get('content_map', {}))} entries: {list(lsp.get('content_map', {}).keys())[:5]}")
        logger.debug(f"First structure unit has {len(structure[0].get('elements', [])) if structure else 0} elements")

        # Iterate through structure units (pages/sections)
        for unit_idx, unit in enumerate(structure):
            unit_type = unit.get("type", "page")
            unit_title = unit.get("title")
            elements = unit.get("elements", [])

            # Add page title if present
            if unit_title and unit_idx == 0:
                heading = doc.add_heading(unit_title, level=0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Render each element
            for element in elements:
                await self._render_element(doc, element, lsp)

            # Add page break between units (except last one)
            if unit_idx < len(structure) - 1:
                doc.add_page_break()

        # Generate artifact ID and save
        artifact_id = f"artifact-word-{uuid4().hex[:8]}"
        output_path = self.output_dir / f"{artifact_id}.docx"

        doc.save(str(output_path))

        logger.info(f"Word document saved: {output_path} (artifact_id: {artifact_id})")

        return artifact_id

    async def _render_element(
        self, doc: Document, element: dict[str, Any], lsp: dict[str, Any]
    ) -> None:
        """
        Render a single layout element.

        Args:
            doc: Document instance
            element: Layout element from LSP
            lsp: Full LSP for content resolution
        """
        element_type = element.get("type", "text")
        content_ref = element.get("content_ref")
        styling = element.get("styling", {})
        gestalt_rules = element.get("gestalt_rules", {})

        # Resolve content from CIP if content_ref is provided
        content_text = self._resolve_content(content_ref, lsp)

        if not content_text or content_text.startswith("[Missing content"):
            logger.warning(f"Skipping element with missing content: content_ref={content_ref}, "
                         f"content_map_keys={list(lsp.get('content_map', {}).keys())[:5]}")
            return

        if element_type == "text":
            # Determine hierarchy level
            hierarchy_level = gestalt_rules.get("hierarchy_level", 4)

            # For long paragraphs, split into sentences for better readability
            # Split on sentence boundaries (period + space, but preserve the period)
            if len(content_text) > 500 and hierarchy_level >= 4:
                # Split into sentences (simple heuristic: period followed by space or end)
                import re
                sentences = re.split(r'(?<=[.!?])\s+', content_text)
                # Filter out empty sentences
                sentences = [s.strip() for s in sentences if s.strip()]
                
                # Add first sentence as paragraph
                if sentences:
                    if hierarchy_level <= 3:
                        paragraph = doc.add_heading(sentences[0], level=hierarchy_level)
                    else:
                        paragraph = doc.add_paragraph(sentences[0])
                    self._apply_styling(paragraph, styling)
                    
                    # Add remaining sentences as separate paragraphs
                    for sentence in sentences[1:]:
                        paragraph = doc.add_paragraph(sentence)
                        self._apply_styling(paragraph, styling)
            else:
                # Add as heading or paragraph (original behavior)
                if hierarchy_level <= 3:
                    paragraph = doc.add_heading(content_text, level=hierarchy_level)
                else:
                    paragraph = doc.add_paragraph(content_text)

                # Apply styling
                self._apply_styling(paragraph, styling)

        elif element_type == "image":
            # For images, content_ref points to image_id
            image_uri = self._resolve_image_uri(content_ref, lsp)
            if image_uri:
                # In production, would download and embed image
                # For v1, add placeholder text
                doc.add_paragraph(f"[Image: {image_uri}]", style="Caption")

    def _resolve_content(self, content_ref: str | None, lsp: dict[str, Any]) -> str | None:
        """
        Resolve content text from content_ref.

        Args:
            content_ref: Block ID to look up
            lsp: LSP containing content

        Returns:
            Content text or None
        """
        if not content_ref:
            return None

        # Look up content in content_map
        content_map = lsp.get("content_map", {})
        return content_map.get(content_ref, f"[Missing content: {content_ref}]")

    def _resolve_image_uri(self, image_id: str | None, lsp: dict[str, Any]) -> str | None:
        """Resolve image URI from image_id."""
        if not image_id:
            return None

        # Look up image URI in content_map
        content_map = lsp.get("content_map", {})
        return content_map.get(image_id, f"[Missing image: {image_id}]")

    def _apply_styling(self, paragraph, styling: dict[str, Any]) -> None:
        """
        Apply styling to paragraph.

        Args:
            paragraph: Paragraph object
            styling: Styling configuration
        """
        # Font styling
        font_config = styling.get("font", {})
        if font_config:
            for run in paragraph.runs:
                font = run.font
                if "size_pt" in font_config:
                    font.size = Pt(font_config["size_pt"])
                if "weight" in font_config and font_config["weight"] == "bold":
                    font.bold = True

        # Color
        color_hex = styling.get("color")
        if color_hex and color_hex.startswith("#"):
            rgb = self._hex_to_rgb(color_hex)
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(*rgb)

        # Alignment
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        alignment = styling.get("alignment", "left")
        paragraph.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

    def _hex_to_rgb(self, hex_color: str) -> tuple[int, int, int]:
        """Convert hex color to RGB tuple."""
        hex_color = hex_color.lstrip("#")
        return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))
